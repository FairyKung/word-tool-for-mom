import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy, io, re, mammoth

st.set_page_config(page_title="แก้คำใน Word", page_icon="📝", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Sarabun', sans-serif !important; }
.stApp { background: linear-gradient(135deg, #f0f4ff 0%, #faf5ff 100%); }
h1 { font-size: 2.2rem !important; font-weight: 700 !important; color: #2d2d6b !important; text-align: center; padding: 0.5rem 0 1.5rem 0; }
h2 { font-size: 1.4rem !important; color: #3b3b8f !important; font-weight: 600 !important; border-left: 5px solid #6c63ff; padding-left: 12px; margin-top: 1.5rem !important; }
[data-testid="stSidebar"] { background: linear-gradient(180deg, #2d2d6b 0%, #3b3b8f 100%) !important; }
[data-testid="stSidebar"] > div { padding-bottom: 2rem; }
[data-testid="stSidebar"] .stTextInput input {
    background: white !important; color: #111111 !important;
    border: 2px solid rgba(255,255,255,0.6) !important; border-radius: 8px !important;
    font-size: 1rem !important; font-family: 'Sarabun', sans-serif !important; padding: 0.5rem 0.8rem !important;
}
[data-testid="stSidebar"] .stTextInput input::placeholder { color: #999 !important; }
[data-testid="stSidebar"] .stTextInput input:focus { border-color: #a78bfa !important; box-shadow: 0 0 0 2px rgba(167,139,250,0.3) !important; }
[data-testid="stSidebar"] label { color: white !important; font-size: 0.95rem !important; font-weight: 600 !important; }
[data-testid="stSidebar"] p, [data-testid="stSidebar"] span, [data-testid="stSidebar"] div, [data-testid="stSidebar"] small { color: white !important; }
.stButton > button { border-radius: 12px !important; font-family: 'Sarabun', sans-serif !important; font-size: 1.05rem !important; font-weight: 600 !important; padding: 0.55rem 1.2rem !important; border: none !important; transition: all 0.2s ease !important; }
.stButton > button[kind="primary"] { background: linear-gradient(135deg, #6c63ff, #4e46d4) !important; color: white !important; width: 100% !important; padding: 0.75rem !important; font-size: 1.15rem !important; box-shadow: 0 4px 14px rgba(108,99,255,0.35) !important; }
.stButton > button[kind="primary"]:hover { transform: translateY(-2px) !important; }
.stButton > button:not([kind="primary"]) { background: rgba(255,255,255,0.2) !important; color: white !important; border: 1.5px solid rgba(255,255,255,0.5) !important; }
[data-testid="stFileUploader"] { background: white !important; border: 2.5px dashed #6c63ff !important; border-radius: 16px !important; padding: 1.5rem !important; }
.paper-wrap { background: #e8eaf6; border-radius: 16px; padding: 24px; box-shadow: inset 0 2px 8px rgba(0,0,0,0.08); }
.paper-page { background: white; border-radius: 8px; padding: 48px 52px; box-shadow: 0 4px 20px rgba(0,0,0,0.12); min-height: 600px; max-width: 820px; margin: 0 auto; font-family: 'Sarabun', sans-serif !important; font-size: 16px; line-height: 1.9; color: #111; }
[data-testid="stDownloadButton"] > button { background: linear-gradient(135deg, #22c55e, #16a34a) !important; color: white !important; font-size: 1.15rem !important; padding: 0.75rem 1.5rem !important; border-radius: 12px !important; font-family: 'Sarabun', sans-serif !important; font-weight: 700 !important; width: 100% !important; box-shadow: 0 4px 14px rgba(34,197,94,0.3) !important; }
[data-testid="stTabs"] [role="tablist"] { background: white; border-radius: 12px; padding: 4px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }
[data-testid="stTabs"] [role="tab"] { font-family: 'Sarabun', sans-serif !important; font-size: 1rem !important; font-weight: 600 !important; border-radius: 10px !important; padding: 0.5rem 1.2rem !important; }
[data-testid="stTabs"] [role="tab"][aria-selected="true"] { background: #6c63ff !important; color: white !important; }
.replace-card { background: rgba(255,255,255,0.1); border: 1px solid rgba(255,255,255,0.2); border-radius: 12px; padding: 14px 14px 8px 14px; margin-bottom: 12px; }
.badge { display: inline-block; background: #818cf8; color: white; font-size: 0.75rem; font-weight: 700; padding: 2px 10px; border-radius: 20px; margin-bottom: 8px; }
.empty-state { text-align: center; padding: 60px 20px; color: #9ca3af; }
.empty-state .icon { font-size: 4rem; }
.empty-state p { font-size: 1.1rem; margin-top: 12px; }
</style>
""", unsafe_allow_html=True)

if 'num_pairs' not in st.session_state:    st.session_state.num_pairs = 3
if 'processed_doc' not in st.session_state: st.session_state.processed_doc = None


# ══════════════════════════════════════════════════════════════════
#  REPLACE ENGINE — ครอบคลุมทุกที่ใน docx
#  ✅ paragraphs ปกติ
#  ✅ tables (skip merged cells)
#  ✅ headers / footers
#  ✅ textboxes (w:txbxContent) ← แก้ปัญหาหน้าที่เปลี่ยนไม่ได้
#  ✅ clone rPr XML เต็ม → ฟอนต์/ขนาดไม่เพี้ยน
#  ✅ รวม run ก่อน regex → คำที่ถูกหั่นข้าม run เปลี่ยนได้
# ══════════════════════════════════════════════════════════════════

def _clone_rpr(run_elem):
    rpr = run_elem.find(qn('w:rPr'))
    return copy.deepcopy(rpr) if rpr is not None else None

def _make_run_with_rpr(paragraph, text, rpr_clone):
    new_run = paragraph.add_run(text)
    existing = new_run._r.find(qn('w:rPr'))
    if existing is not None: new_run._r.remove(existing)
    if rpr_clone is not None: new_run._r.insert(0, copy.deepcopy(rpr_clone))
    return new_run

def _replace_in_paragraph(paragraph, old_text, new_text):
    runs = paragraph.runs
    if not runs: return 0
    rpr_list = [_clone_rpr(r._r) for r in runs]
    full_text, char_to_run = "", []
    for r_idx, run in enumerate(runs):
        for _ in run.text: char_to_run.append(r_idx)
        full_text += run.text
    pattern = re.escape(old_text).replace(r'\ ', r'\s+')
    matches = list(re.finditer(pattern, full_text))
    if not matches: return 0
    segments, cursor = [], 0
    for m in matches:
        for i in range(cursor, m.start()):
            r_idx = char_to_run[i]
            if segments and segments[-1][1] == r_idx: segments[-1] = (segments[-1][0] + full_text[i], r_idx)
            else: segments.append((full_text[i], r_idx))
        first_r = char_to_run[m.start()] if m.start() < len(char_to_run) else 0
        if new_text:
            if segments and segments[-1][1] == first_r: segments[-1] = (segments[-1][0] + new_text, first_r)
            else: segments.append((new_text, first_r))
        cursor = m.end()
    for i in range(cursor, len(full_text)):
        r_idx = char_to_run[i] if i < len(char_to_run) else len(runs)-1
        if segments and segments[-1][1] == r_idx: segments[-1] = (segments[-1][0] + full_text[i], r_idx)
        else: segments.append((full_text[i], r_idx))
    p_elem = paragraph._p
    for run in runs: p_elem.remove(run._r)
    for seg_text, r_idx in segments:
        if seg_text: _make_run_with_rpr(paragraph, seg_text, rpr_list[r_idx])
    return len(matches)

def _replace_in_paragraphs(paragraphs, old, new):
    return sum(_replace_in_paragraph(p, old, new) for p in paragraphs)

def _replace_in_table(table, old, new):
    count = 0
    for row in table.rows:
        seen_tc = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc: continue
            seen_tc.add(tc_id)
            count += _replace_in_paragraphs(cell.paragraphs, old, new)
    return count

def _replace_in_textboxes(doc, old, new):
    """แก้คำใน textbox (w:txbxContent) ซึ่ง python-docx ไม่ได้ expose ออกมา"""
    from docx.text.paragraph import Paragraph as DocxParagraph
    count = 0
    body = doc.element.body
    for txbx in body.findall('.//' + qn('w:txbxContent')):
        for p_elem in txbx.findall('.//' + qn('w:p')):
            # ห่อ p element เป็น Paragraph object ของ python-docx
            p_obj = DocxParagraph(p_elem, doc)
            count += _replace_in_paragraph(p_obj, old, new)
    return count

def process_document(file_bytes, replace_list):
    doc = Document(io.BytesIO(file_bytes))
    total = 0
    for old, new in replace_list:
        if not old.strip(): continue
        # 1. paragraphs ปกติ
        total += _replace_in_paragraphs(doc.paragraphs, old, new)
        # 2. tables (skip merged cells)
        for table in doc.tables:
            total += _replace_in_table(table, old, new)
        # 3. headers / footers
        for section in doc.sections:
            for hf in [section.header, section.footer,
                       section.first_page_header, section.first_page_footer,
                       section.even_page_header, section.even_page_footer]:
                if hf: total += _replace_in_paragraphs(hf.paragraphs, old, new)
        # 4. textboxes ← ใหม่!
        total += _replace_in_textboxes(doc, old, new)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), total


# ══════════════════════════════════════════════════════════════════
#  UI
# ══════════════════════════════════════════════════════════════════
st.markdown("# 📝 โปรแกรมแก้คำใน Word")
st.markdown("<p style='text-align:center;color:#6b7280;font-size:1.05rem;margin-top:-12px;margin-bottom:24px;'>อัปโหลดไฟล์ · ใส่คำที่ต้องการเปลี่ยน · ดาวน์โหลดไฟล์ใหม่ได้เลย</p>", unsafe_allow_html=True)

uploaded_file = st.file_uploader("📂  เลือกหรือลากไฟล์ Word มาวางที่นี่", type=["docx"])

with st.sidebar:
    st.markdown("## 🔄 รายการคำที่ต้องการเปลี่ยน")
    st.markdown("<p style='font-size:0.9rem;opacity:0.8;margin-top:-8px;margin-bottom:12px;'>ใส่คำเดิมและคำใหม่ที่ต้องการแทน</p>", unsafe_allow_html=True)
    st.markdown("---")
    replace_list = []
    for i in range(st.session_state.num_pairs):
        st.markdown(f'<div class="replace-card"><span class="badge">คู่ที่ {i+1}</span>', unsafe_allow_html=True)
        old = st.text_input("คำเดิม", key=f"old_{i}", placeholder="เช่น 1 เมษายน 2568")
        new = st.text_input("คำใหม่", key=f"new_{i}", placeholder="เช่น 1 พฤษภาคม 2568")
        st.markdown('</div>', unsafe_allow_html=True)
        if old.strip(): replace_list.append((old.strip(), new.strip()))
    c1, c2 = st.columns(2)
    with c1:
        if st.button("➕ เพิ่มช่อง"):
            st.session_state.num_pairs += 1; st.rerun()
    with c2:
        if st.button("🗑️ ล้างทั้งหมด"):
            st.session_state.num_pairs = 3
            for k in list(st.session_state.keys()):
                if k.startswith("old_") or k.startswith("new_"): del st.session_state[k]
            st.session_state.processed_doc = None; st.rerun()
    st.markdown("---")
    process_btn = st.button("🚀 เริ่มเปลี่ยนคำ", type="primary",
                            disabled=(uploaded_file is None or len(replace_list) == 0))
    if uploaded_file is None:  st.caption("⬆️ กรุณาอัปโหลดไฟล์ก่อน")
    elif not replace_list:     st.caption("⬆️ กรุณาใส่คำที่ต้องการเปลี่ยนก่อน")
    if replace_list:
        st.markdown("---")
        st.markdown("**📋 สรุปรายการที่จะเปลี่ยน:**")
        for old, new in replace_list:
            st.markdown(f"• `{old}` → {'`'+new+'`' if new else '_ลบออก_'}")

if uploaded_file is not None:
    file_content = uploaded_file.getvalue()
    if process_btn and replace_list:
        with st.spinner("⏳ กำลังแก้ไขไฟล์..."):
            result_bytes, count = process_document(file_content, replace_list)
            st.session_state.processed_doc = result_bytes
        st.success(f"✅ แก้ไขเรียบร้อย! พบและเปลี่ยนใน {count} ตำแหน่ง")
    if st.session_state.processed_doc:
        st.download_button("📥  ดาวน์โหลดไฟล์ที่แก้เสร็จแล้ว",
                           data=st.session_state.processed_doc,
                           file_name=f"แก้แล้ว_{uploaded_file.name}",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.markdown("---")
        tab1, tab2 = st.tabs(["📄 ต้นฉบับ", "✨ หลังแก้ไข"])
    else:
        tab1, = st.tabs(["📄 ต้นฉบับ"])
        tab2 = None
    with tab1:
        st.markdown("## 📄 ไฟล์ต้นฉบับ")
        try:
            html = mammoth.convert_to_html(io.BytesIO(file_content)).value
            st.markdown(f'<div class="paper-wrap"><div class="paper-page">{html}</div></div>', unsafe_allow_html=True)
        except Exception as e: st.error(f"Preview ไม่ได้: {e}")
    if tab2 and st.session_state.processed_doc:
        with tab2:
            st.markdown("## ✨ ไฟล์หลังแก้ไข")
            if replace_list:
                st.info("🔄 " + " · ".join([f"`{o}` → `{n}`" if n else f"`{o}` → _(ลบ)_" for o,n in replace_list]))
            try:
                html = mammoth.convert_to_html(io.BytesIO(st.session_state.processed_doc)).value
                st.markdown(f'<div class="paper-wrap"><div class="paper-page">{html}</div></div>', unsafe_allow_html=True)
            except Exception as e: st.error(f"Preview ไม่ได้: {e}")
else:
    st.markdown('<div class="empty-state"><div class="icon">📂</div><p>ยังไม่ได้เลือกไฟล์<br>กรุณาอัปโหลดไฟล์ Word (.docx) ด้านบนก่อนนะคะ</p></div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("## 📖 วิธีใช้งาน")
    c1, c2, c3 = st.columns(3)
    with c1: st.markdown("### 1️⃣ อัปโหลดไฟล์\nคลิกที่กล่องด้านบน แล้วเลือกไฟล์ Word ที่ต้องการแก้ไข")
    with c2: st.markdown("### 2️⃣ ใส่คำที่ต้องการเปลี่ยน\nใส่คำเดิมและคำใหม่ในแถบด้านซ้าย สามารถเพิ่มได้หลายคู่")
    with c3: st.markdown("### 3️⃣ ดาวน์โหลด\nกด **เริ่มเปลี่ยนคำ** แล้วดาวน์โหลดไฟล์ที่แก้เสร็จแล้ว")